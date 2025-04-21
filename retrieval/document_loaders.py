"""
Document Loaders Module
Handles loading and preprocessing various document formats (PDF, DOCX, HTML, etc.)
"""
import json
import logging
import os
import re
import sys
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Union

# Add project root to path to import config
sys.path.append(str(Path(__file__).parent.parent.absolute()))
import config

# Set up logging
logging.basicConfig(**config.LOGGING_CONFIG)
logger = logging.getLogger(__name__)


class BaseDocumentLoader:
    """Base class for document loaders."""
    
    def __init__(self, remove_empty_lines: bool = True, remove_multiple_whitespaces: bool = True):
        """
        Initialize base document loader.
        
        Args:
            remove_empty_lines: Whether to remove empty lines
            remove_multiple_whitespaces: Whether to collapse multiple whitespaces
        """
        self.remove_empty_lines = remove_empty_lines
        self.remove_multiple_whitespaces = remove_multiple_whitespaces
    
    def load(self, file_path: str) -> Optional[Dict]:
        """
        Load document from file.
        
        Args:
            file_path: Path to file
            
        Returns:
            Document dictionary or None if loading fails
        """
        raise NotImplementedError("Subclasses must implement load()")
    
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text.
        
        Args:
            text: Text to clean
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Replace multiple newlines with a single newline
        if self.remove_empty_lines:
            text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Replace multiple whitespaces with a single space
        if self.remove_multiple_whitespaces:
            text = re.sub(r'\s+', ' ', text)
        
        return text.strip()


class TextLoader(BaseDocumentLoader):
    """Loader for plain text files."""
    
    def load(self, file_path: str) -> Optional[Dict]:
        """
        Load document from a text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Document dictionary or None if loading fails
        """
        try:
            path = Path(file_path)
            
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            text = self.clean_text(text)
            
            document = {
                "id": path.stem,
                "title": path.name,
                "text": text,
                "source": str(path),
                "metadata": {
                    "file_type": "txt",
                    "file_size": os.path.getsize(path)
                }
            }
            
            logger.info(f"Loaded text document: {path.name}")
            return document
            
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            return None


class PDFLoader(BaseDocumentLoader):
    """Loader for PDF files."""
    
    def load(self, file_path: str) -> Optional[Dict]:
        """
        Load document from a PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Document dictionary or None if loading fails
        """
        try:
            import pypdf
            
            path = Path(file_path)
            
            # Extract text from PDF
            text_parts = []
            metadata = {}
            
            with open(path, 'rb') as f:
                pdf = pypdf.PdfReader(f)
                
                # Extract metadata
                if pdf.metadata:
                    for key, value in pdf.metadata.items():
                        if key and value:
                            # Convert PDF metadata key format /Title to title
                            if key.startswith('/'):
                                key = key[1:].lower()
                            metadata[key] = str(value)
                
                # Extract text from each page
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            # Join all text parts
            full_text = "\n\n".join(text_parts)
            full_text = self.clean_text(full_text)
            
            # Determine title from metadata or filename
            title = metadata.get("title", path.stem)
            
            document = {
                "id": path.stem,
                "title": title,
                "text": full_text,
                "source": str(path),
                "metadata": {
                    "file_type": "pdf",
                    "file_size": os.path.getsize(path),
                    "pages": len(pdf.pages),
                    **metadata
                }
            }
            
            logger.info(f"Loaded PDF document: {path.name} ({len(pdf.pages)} pages)")
            return document
            
        except ImportError:
            logger.error("pypdf not installed. Please install it with 'pip install pypdf'")
            return None
        except Exception as e:
            logger.error(f"Error loading PDF file {file_path}: {e}")
            return None


class DocxLoader(BaseDocumentLoader):
    """Loader for DOCX files."""
    
    def load(self, file_path: str) -> Optional[Dict]:
        """
        Load document from a DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Document dictionary or None if loading fails
        """
        try:
            import docx
            
            path = Path(file_path)
            
            # Load DOCX document
            doc = docx.Document(path)
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)
            
            # Get text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)
            
            # Join all text parts
            full_text = "\n\n".join(text_parts)
            full_text = self.clean_text(full_text)
            
            # Get core properties
            metadata = {}
            try:
                core_properties = doc.core_properties
                if core_properties.title:
                    metadata["title"] = core_properties.title
                if core_properties.author:
                    metadata["author"] = core_properties.author
                if core_properties.created:
                    metadata["created"] = str(core_properties.created)
                if core_properties.modified:
                    metadata["modified"] = str(core_properties.modified)
            except:
                pass
            
            # Determine title from metadata or filename
            title = metadata.get("title", path.stem)
            
            document = {
                "id": path.stem,
                "title": title,
                "text": full_text,
                "source": str(path),
                "metadata": {
                    "file_type": "docx",
                    "file_size": os.path.getsize(path),
                    "paragraphs": len(doc.paragraphs),
                    **metadata
                }
            }
            
            logger.info(f"Loaded DOCX document: {path.name} ({len(doc.paragraphs)} paragraphs)")
            return document
            
        except ImportError:
            logger.error("python-docx not installed. Please install it with 'pip install python-docx'")
            return None
        except Exception as e:
            logger.error(f"Error loading DOCX file {file_path}: {e}")
            return None


class HTMLLoader(BaseDocumentLoader):
    """Loader for HTML files."""
    
    def load(self, file_path: str) -> Optional[Dict]:
        """
        Load document from an HTML file.
        
        Args:
            file_path: Path to HTML file
            
        Returns:
            Document dictionary or None if loading fails
        """
        try:
            from bs4 import BeautifulSoup
            
            path = Path(file_path)
            
            with open(path, 'r', encoding='utf-8') as f:
                html = f.read()
            
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extract title
            title_tag = soup.find('title')
            title = title_tag.get_text() if title_tag else path.stem
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            # Get text
            text = soup.get_text()
            text = self.clean_text(text)
            
            # Extract metadata
            metadata = {}
            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                name = tag.get('name', tag.get('property', ''))
                content = tag.get('content', '')
                if name and content:
                    metadata[name] = content
            
            document = {
                "id": path.stem,
                "title": title,
                "text": text,
                "source": str(path),
                "metadata": {
                    "file_type": "html",
                    "file_size": os.path.getsize(path),
                    **metadata
                }
            }
            
            logger.info(f"Loaded HTML document: {path.name}")
            return document
            
        except ImportError:
            logger.error("beautifulsoup4 not installed. Please install it with 'pip install beautifulsoup4'")
            return None
        except Exception as e:
            logger.error(f"Error loading HTML file {file_path}: {e}")
            return None


class CSVLoader(BaseDocumentLoader):
    """Loader for CSV files."""
    
    def __init__(
        self,
        content_columns: Optional[List[str]] = None,
        header_row: bool = True,
        delimiter: str = ',',
        **kwargs
    ):
        """
        Initialize CSV loader.
        
        Args:
            content_columns: Columns to include in document text
            header_row: Whether the CSV has a header row
            delimiter: CSV delimiter
        """
        super().__init__(**kwargs)
        self.content_columns = content_columns
        self.header_row = header_row
        self.delimiter = delimiter
    
    def load(self, file_path: str) -> Optional[Dict]:
        """
        Load document from a CSV file.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Document dictionary or None if loading fails
        """
        try:
            import csv
            
            path = Path(file_path)
            
            with open(path, 'r', encoding='utf-8') as f:
                # Read CSV file
                reader = csv.reader(f, delimiter=self.delimiter)
                rows = list(reader)
            
            if not rows:
                logger.warning(f"Empty CSV file: {path.name}")
                return None
            
            # Get header row if present
            headers = rows[0] if self.header_row else None
            data_rows = rows[1:] if self.header_row else rows
            
            # Determine which columns to include
            if self.content_columns and headers:
                indices = [headers.index(col) for col in self.content_columns if col in headers]
            elif headers:
                indices = list(range(len(headers)))
            else:
                indices = list(range(len(rows[0])))
            
            # Build text from selected columns
            text_parts = []
            
            for row in data_rows:
                # Skip rows that are too short
                if len(row) <= max(indices, default=0):
                    continue
                
                # Extract values from selected columns
                if headers:
                    # Include column headers in output
                    row_parts = [f"{headers[i]}: {row[i]}" for i in indices]
                else:
                    row_parts = [row[i] for i in indices]
                
                text_parts.append(", ".join(row_parts))
            
            # Join all text parts
            full_text = "\n".join(text_parts)
            full_text = self.clean_text(full_text)
            
            document = {
                "id": path.stem,
                "title": path.name,
                "text": full_text,
                "source": str(path),
                "metadata": {
                    "file_type": "csv",
                    "file_size": os.path.getsize(path),
                    "rows": len(data_rows),
                    "columns": len(headers) if headers else len(rows[0]) if rows else 0
                }
            }
            
            logger.info(f"Loaded CSV document: {path.name} ({len(data_rows)} rows)")
            return document
            
        except Exception as e:
            logger.error(f"Error loading CSV file {file_path}: {e}")
            return None


class WebPageLoader(BaseDocumentLoader):
    """Loader for web pages."""
    
    def load(self, url: str) -> Optional[Dict]:
        """
        Load document from a web page.
        
        Args:
            url: URL of web page
            
        Returns:
            Document dictionary or None if loading fails
        """
        try:
            import requests
            from bs4 import BeautifulSoup
            
            # Request web page
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            
            # Parse HTML
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Extract title
            title_tag = soup.find('title')
            title = title_tag.get_text() if title_tag else url
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            # Get text
            text = soup.get_text()
            text = self.clean_text(text)
            
            # Extract metadata
            metadata = {
                "url": url,
                "status_code": response.status_code
            }
            
            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                name = tag.get('name', tag.get('property', ''))
                content = tag.get('content', '')
                if name and content:
                    metadata[name] = content
            
            # Generate an ID from the URL
            import hashlib
            url_hash = hashlib.md5(url.encode()).hexdigest()
            
            document = {
                "id": f"webpage_{url_hash}",
                "title": title,
                "text": text,
                "source": url,
                "metadata": metadata
            }
            
            logger.info(f"Loaded web page: {url}")
            return document
            
        except ImportError:
            logger.error("Required libraries not installed. Please install with 'pip install requests beautifulsoup4'")
            return None
        except Exception as e:
            logger.error(f"Error loading web page {url}: {e}")
            return None


def load_document(file_path: str) -> Optional[Dict]:
    """
    Load a document from a file based on its extension.
    
    Args:
        file_path: Path to document
        
    Returns:
        Document dictionary or None if loading fails
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    # Select loader based on file extension
    if suffix == '.txt':
        loader = TextLoader()
    elif suffix == '.pdf':
        loader = PDFLoader()
    elif suffix == '.docx':
        loader = DocxLoader()
    elif suffix == '.html' or suffix == '.htm':
        loader = HTMLLoader()
    elif suffix == '.csv':
        loader = CSVLoader()
    elif suffix.startswith('.'):
        # Handle any other file extension by treating as text
        logger.warning(f"No specific loader for {suffix}, using text loader")
        loader = TextLoader()
    else:
        logger.error(f"Unsupported file type: {file_path}")
        return None
    
    return loader.load(file_path)


def load_documents_from_directory(
    directory_path: str,
    recursive: bool = True,
    extensions: Optional[List[str]] = None
) -> List[Dict]:
    """
    Load documents from a directory.
    
    Args:
        directory_path: Path to directory
        recursive: Whether to search recursively
        extensions: List of file extensions to include (e.g., ['.pdf', '.txt'])
        
    Returns:
        List of document dictionaries
    """
    if extensions is None:
        extensions = ['.txt', '.pdf', '.docx', '.html', '.htm', '.csv']
    
    path = Path(directory_path)
    documents = []
    
    # Get all matching files
    all_files = []
    if recursive:
        for ext in extensions:
            all_files.extend(path.glob(f"**/*{ext}"))
    else:
        for ext in extensions:
            all_files.extend(path.glob(f"*{ext}"))
    
    # Load each file
    for file_path in all_files:
        doc = load_document(str(file_path))
        if doc:
            documents.append(doc)
    
    logger.info(f"Loaded {len(documents)} documents from directory: {directory_path}")
    return documents


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Load and preprocess documents")
    parser.add_argument("--input", type=str, required=True, help="Input file or directory")
    parser.add_argument("--output", type=str, required=True, help="Output JSON file")
    parser.add_argument("--recursive", action="store_true", help="Search directory recursively")
    
    args = parser.parse_args()
    
    # Process input
    path = Path(args.input)
    
    if path.is_file():
        # Load single file
        doc = load_document(str(path))
        documents = [doc] if doc else []
    elif path.is_dir():
        # Load directory
        documents = load_documents_from_directory(
            str(path),
            recursive=args.recursive
        )
    else:
        logger.error(f"Input path does not exist: {path}")
        sys.exit(1)
    
    if not documents:
        logger.error(f"No documents loaded from {path}")
        sys.exit(1)
    
    # Save to output file
    output_path = Path(args.output)
    output_path.parent.mkdir(exist_ok=True, parents=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(documents, f, indent=2)
    
    logger.info(f"Saved {len(documents)} documents to {output_path}")
